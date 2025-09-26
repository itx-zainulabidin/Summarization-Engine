import os
from pathlib import Path
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)


def export_txt(summary: str, doc_id: str, mode: str) -> str:
    """
    Export summary as plain text (.txt).

    Args:
        summary (str): The summary text to save
        doc_id (str): Original document identifier
        mode (str): Summarization mode (tldr, short, extended)

    Returns:
        str: Path to the exported file
    """
    
    filename = EXPORT_DIR / f"{doc_id}_{mode}.txt"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(summary)
    return str(filename)


def export_docx(summary: str, doc_id: str, mode: str) -> str:
    """
    Export summary as Microsoft Word (.docx).

    Args:
        summary (str): The summary text to save
        doc_id (str): Original document identifier
        mode (str): Summarization mode (tldr, short, extended)

    Returns:
        str: Path to the exported file
    """

    filename = EXPORT_DIR / f"{doc_id}_{mode}.docx"
    doc = Document()
    doc.add_heading(f"Summary ({mode})", 0)
    doc.add_paragraph(summary)
    doc.save(filename)
    return str(filename)


def export_pdf(summary: str, doc_id: str, mode: str) -> str:
    """
    Export summary as a PDF file.

    Args:
        summary (str): The summary text to save
        doc_id (str): Original document identifier
        mode (str): Summarization mode (tldr, short, extended)

    Returns:
        str: Path to the exported file
    """
    
    filename = EXPORT_DIR / f"{doc_id}_{mode}.pdf"
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(filename))
    elements = [Paragraph(f"Summary ({mode})", styles["Title"]),
                Paragraph(summary, styles["Normal"])]
    doc.build(elements)
    return str(filename)


def export_summary(summary: str, doc_id: str, mode: str, export_format: str) -> str:
    """
    Route export to the correct format (txt, docx, pdf).

    Args:
        summary (str): The summary text to save
        doc_id (str): Original document identifier
        mode (str): Summarization mode (tldr, short, extended)
        export_format (str): File format to export ("txt", "docx", "pdf")

    Returns:
        str: Path to the exported file

    Raises:
        ValueError: If unsupported export format is requested
    """
    
    if export_format == "txt":
        return export_txt(summary, doc_id, mode)
    elif export_format == "docx":
        return export_docx(summary, doc_id, mode)
    elif export_format == "pdf":
        return export_pdf(summary, doc_id, mode)
    else:
        raise ValueError(f"Unsupported export format: {export_format}")
